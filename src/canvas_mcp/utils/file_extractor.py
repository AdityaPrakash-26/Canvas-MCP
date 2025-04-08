"""
File Content Extraction Utilities

This module provides utilities for extracting content from different file types,
such as DOCX, PDF, HTML/URL, etc. It's used for processing syllabi and other documents
in the Canvas MCP server.
"""

import logging
import os
import tempfile
from typing import Any

import pdfplumber
import requests
from bs4 import BeautifulSoup

# Configure logging
logger = logging.getLogger(__name__)


def download_file(
    url: str, suffix: str = None, timeout: int = 30
) -> tuple[str, str, Exception | None]:
    """
    Download a file from a URL to a temporary local file.

    Args:
        url: URL of the file to download
        suffix: Optional suffix for the temporary file (e.g., '.pdf')
        timeout: Request timeout in seconds

    Returns:
        Tuple of (file_path, content_type, error)
        - file_path: Path to the downloaded file (None if download failed)
        - content_type: Content type of the downloaded file
        - error: Exception if an error occurred, None otherwise
    """
    try:
        # Default suffix based on URL if not provided
        if suffix is None:
            if url.lower().endswith(".pdf"):
                suffix = ".pdf"
            elif url.lower().endswith(".docx"):
                suffix = ".docx"
            elif url.lower().endswith(".doc"):
                suffix = ".doc"
            else:
                suffix = ".tmp"

        # Download the file
        response = requests.get(url, stream=True, timeout=timeout)
        response.raise_for_status()  # Raise exception for bad responses

        # Get content type
        content_type = response.headers.get("Content-Type", "").lower()

        # Create a temporary file to save the content
        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as temp_file:
            temp_path = temp_file.name
            # Write the content to the temporary file
            for chunk in response.iter_content(chunk_size=8192):
                temp_file.write(chunk)

        return temp_path, content_type, None

    except Exception as e:
        logger.error(f"Error downloading file from {url}: {e}")
        return None, None, e


def extract_text_from_url(url: str, timeout: int = 30) -> str | None:
    """
    Fetch content from a URL and extract its text content.

    Args:
        url: URL to fetch content from
        timeout: Request timeout in seconds

    Returns:
        Extracted text as a string or None if extraction failed
    """
    try:
        # Download the content
        response = requests.get(url, timeout=timeout)
        response.raise_for_status()  # Raise exception for bad responses

        content_type = response.headers.get("Content-Type", "").lower()

        # Handle based on content type
        if "application/pdf" in content_type:
            # It's a PDF, use the PDF extraction function
            return extract_text_from_pdf_url(url, max_pages=50)
        elif "text/html" in content_type:
            # It's an HTML page, extract text
            soup = BeautifulSoup(response.text, "html.parser")

            # Remove script and style elements
            for script_or_style in soup(["script", "style"]):
                script_or_style.decompose()

            # Get text
            text = soup.get_text(separator="\n")

            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = "\n".join(chunk for chunk in chunks if chunk)

            return text
        else:
            # Just return the raw content as text
            return response.text

    except Exception as e:
        logger.error(f"Error extracting text from URL {url}: {e}")
        return None


def extract_text_from_pdf_url(url: str, max_pages: int = 50) -> str | None:
    """
    Download a PDF from a URL and extract its text content.

    Args:
        url: URL of the PDF file
        max_pages: Maximum number of pages to extract (default: 50)

    Returns:
        Extracted text as a string or None if extraction failed
    """
    try:
        # Download the file
        temp_path, content_type, error = download_file(url, suffix=".pdf")

        if error:
            return None

        # Check if it's actually a PDF
        if (
            content_type
            and "application/pdf" not in content_type
            and not url.lower().endswith(".pdf")
        ):
            logger.warning(
                f"URL {url} does not appear to be a PDF (Content-Type: {content_type})"
            )
            # Try to extract as a regular URL
            os.unlink(temp_path)  # Clean up the temporary file
            return extract_text_from_url(url)

        # Extract text from the downloaded PDF
        text = extract_text_from_pdf_file(temp_path, max_pages)

        # Clean up the temporary file
        os.unlink(temp_path)

        return text

    except Exception as e:
        logger.error(f"Error extracting text from PDF URL {url}: {e}")
        return None


def extract_text_from_pdf_file(file_path: str, max_pages: int = 50) -> str | None:
    """
    Extract text content from a local PDF file.

    Args:
        file_path: Path to the PDF file
        max_pages: Maximum number of pages to extract (default: 50)

    Returns:
        Extracted text as a string or None if extraction failed
    """
    try:
        text_content = []

        with pdfplumber.open(file_path) as pdf:
            # Limit to max_pages or the actual page count, whichever is smaller
            pages_to_extract = min(len(pdf.pages), max_pages)

            for i in range(pages_to_extract):
                page = pdf.pages[i]
                text = page.extract_text() or ""
                if text:
                    text_content.append(text)

                # Also extract tables as text if available
                tables = page.extract_tables()
                if tables:
                    for table in tables:
                        table_text = "\n".join(
                            [" | ".join([cell or "" for cell in row]) for row in table]
                        )
                        text_content.append(f"\nTable:\n{table_text}\n")

        return "\n\n".join(text_content)

    except Exception as e:
        logger.error(f"Error extracting text from PDF file {file_path}: {e}")
        return None


def extract_text_from_docx_url(url: str) -> str | None:
    """
    Download a DOCX file from a URL and extract its text content.

    Args:
        url: URL of the DOCX file

    Returns:
        Extracted text as a string or None if extraction failed
    """
    try:
        # Check if python-docx is installed
        try:
            import docx
        except ImportError:
            logger.error("python-docx is not installed. Cannot extract text from DOCX.")
            return None

        # Download the file
        temp_path, content_type, error = download_file(url, suffix=".docx")

        if error:
            return None

        # Check if it's actually a DOCX
        docx_types = [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/docx",
            "application/msword",
        ]

        is_docx = (
            content_type and any(type in content_type for type in docx_types)
        ) or url.lower().endswith((".docx", ".doc"))

        if not is_docx:
            logger.warning(
                f"URL {url} does not appear to be a DOCX (Content-Type: {content_type})"
            )
            # Try to extract as a regular URL
            os.unlink(temp_path)  # Clean up the temporary file
            return extract_text_from_url(url)

        # Extract text from the downloaded DOCX
        text = extract_text_from_docx_file(temp_path)

        # Clean up the temporary file
        os.unlink(temp_path)

        return text

    except Exception as e:
        logger.error(f"Error extracting text from DOCX URL {url}: {e}")
        return None


def extract_text_from_docx_file(file_path: str) -> str | None:
    """
    Extract text content from a local DOCX file.

    Args:
        file_path: Path to the DOCX file

    Returns:
        Extracted text as a string or None if extraction failed
    """
    try:
        # Check if python-docx is installed
        try:
            import docx
        except ImportError:
            logger.error("python-docx is not installed. Cannot extract text from DOCX.")
            return None

        # Open the document
        doc = docx.Document(file_path)

        # Extract text from paragraphs
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Extract text from tables
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(" | ".join(row_data))
            tables.append("\nTable:\n" + "\n".join(table_data))

        # Combine all extracted text
        all_text = paragraphs + tables
        return "\n\n".join(all_text)

    except Exception as e:
        logger.error(f"Error extracting text from DOCX file {file_path}: {e}")
        return None


def extract_text_from_file(source: str, file_type: str = None) -> dict[str, Any]:
    """
    Extract text from a file based on its type.

    Args:
        source: Path or URL to the file
        file_type: Type of file ('pdf', 'docx', 'url', etc.) or None to auto-detect

    Returns:
        Dictionary with extracted content and metadata:
            {
                "success": bool,  # Whether extraction was successful
                "file_type": str,  # Detected or provided file type
                "text": str,      # Extracted text (None if extraction failed)
                "error": str      # Error message (None if successful)
            }
    """
    if not source:
        return {
            "success": False,
            "error": "No source provided",
            "file_type": None,
            "text": None,
        }

    # Auto-detect file type if not provided
    if not file_type:
        # Check if it's a URL or local file
        is_url = source.lower().startswith(("http://", "https://"))

        # Get file extension
        file_extension = None
        if "." in source.split("/")[-1]:
            file_extension = source.split(".")[-1].lower()

        # Determine file type
        if file_extension == "pdf":
            file_type = "pdf"
        elif file_extension in ["docx", "doc"]:
            file_type = "docx"
        elif is_url:
            file_type = "url"
        else:
            # Try to guess by content for local files
            if os.path.isfile(source):
                with open(source, "rb") as f:
                    header = f.read(8)  # Read first 8 bytes
                    # %PDF magic number
                    if header.startswith(b"%PDF"):
                        file_type = "pdf"
                    elif header.startswith(b"PK\x03\x04"):  # ZIP file (could be DOCX)
                        file_type = "docx"
                    else:
                        file_type = "unknown"
            else:
                file_type = "unknown"

    text = None
    success = False
    error = None

    try:
        # Extract based on file type
        if file_type == "pdf":
            if source.lower().startswith(("http://", "https://")):
                text = extract_text_from_pdf_url(source)
            else:
                text = extract_text_from_pdf_file(source)
        elif file_type == "docx":
            if source.lower().startswith(("http://", "https://")):
                text = extract_text_from_docx_url(source)
            else:
                text = extract_text_from_docx_file(source)
        elif file_type == "url":
            text = extract_text_from_url(source)
        else:
            error = f"Unsupported file type: {file_type}"

        success = text is not None
        if not success and not error:
            error = f"Failed to extract text from {file_type} file"

    except Exception as e:
        success = False
        error = f"Error extracting text: {str(e)}"

    return {
        "success": success,
        "file_type": file_type,
        "text": text,
        "error": error,
        "source": source,
    }
